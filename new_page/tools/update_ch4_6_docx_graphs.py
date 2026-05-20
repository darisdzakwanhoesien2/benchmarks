from __future__ import annotations

from pathlib import Path
import shutil
import textwrap
from typing import Iterable

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT.parent / "pages" / "thesis_ch4_6_structure_benchmarks.docx"
OUTPUT_DOCX = ROOT.parent / "pages" / "thesis_ch4_6_structure_benchmarks_streamlit_graphs.docx"
GRAPH_DIR = ROOT / "results" / "docx_graph_attachments"
VIS = ROOT / "results" / "visualizations"
REV = ROOT / "results" / "revision_analysis"


def load_csv(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path).fillna("")


def font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica.ttc",
        "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except Exception:
            continue
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, width: int, fnt: ImageFont.ImageFont) -> list[str]:
    words = str(text).split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= width or not current:
            current = trial
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_bar_chart(
    path: Path,
    title: str,
    rows: Iterable[tuple[str, float]],
    *,
    subtitle: str = "",
    value_suffix: str = "",
    color: str = "#2f6f73",
) -> Path:
    rows = [(str(label), float(value)) for label, value in rows if pd.notna(value)]
    rows = rows[:14]
    w, h = 1800, max(720, 170 + len(rows) * 78)
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(46, True)
    subtitle_font = font(25)
    label_font = font(25)
    small_font = font(23)
    draw.rectangle((0, 0, w, 104), fill="#eef6f4")
    draw.text((52, 28), title, fill="#173f42", font=title_font)
    if subtitle:
        draw.text((54, 112), subtitle, fill="#5b6472", font=subtitle_font)
    max_value = max([value for _, value in rows] or [1])
    x0, x1 = 520, w - 160
    y = 185
    bar_h = 38
    for label, value in rows:
        for idx, line in enumerate(wrap_text(draw, label, 430, label_font)[:2]):
            draw.text((54, y - 3 + idx * 27), line, fill="#1f2937", font=label_font)
        bar_w = int((x1 - x0) * (value / max_value)) if max_value else 0
        draw.rounded_rectangle((x0, y, x0 + bar_w, y + bar_h), radius=8, fill=color)
        draw.line((x0, y + bar_h + 12, x1, y + bar_h + 12), fill="#e5e7eb", width=2)
        value_text = f"{value:.3f}{value_suffix}" if abs(value) <= 1 and not value_suffix else f"{value:,.0f}{value_suffix}"
        draw.text((x0 + bar_w + 18, y + 3), value_text, fill="#111827", font=small_font)
        y += 78
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)
    return path


def create_generated_graphs() -> list[tuple[Path, str, str]]:
    GRAPH_DIR.mkdir(parents=True, exist_ok=True)
    attachments: list[tuple[Path, str, str]] = []

    model = load_csv(REV / "model_stability_summary.csv")
    if not model.empty and {"model", "json_parse_success_rate"}.issubset(model.columns):
        plot = model.copy()
        plot["json_parse_success_rate"] = pd.to_numeric(plot["json_parse_success_rate"], errors="coerce")
        plot = plot.sort_values("json_parse_success_rate", ascending=False)
        p = draw_bar_chart(
            GRAPH_DIR / "docx_model_parse_success.png",
            "Model Parse Success Benchmark",
            zip(plot["model"], plot["json_parse_success_rate"]),
            subtitle="Source: pages/6_1 and pages/6_3 model_stability_chart",
            color="#395b91",
        )
        attachments.append((p, "Figure A.10 - Model parse success benchmark", "Compares model configurations by JSON parse success rate for RQ6 stability."))

    prompt = load_csv(REV / "prompt_stability_summary.csv")
    if not prompt.empty and {"prompt", "missing_tone_rate"}.issubset(prompt.columns):
        plot = prompt.copy()
        plot["missing_tone_rate"] = pd.to_numeric(plot["missing_tone_rate"], errors="coerce")
        plot = plot.sort_values("missing_tone_rate", ascending=True)
        p = draw_bar_chart(
            GRAPH_DIR / "docx_prompt_missing_tone_rate.png",
            "Prompt Missing-Tone Rate Benchmark",
            zip(plot["prompt"], plot["missing_tone_rate"]),
            subtitle="Lower is better; source: pages/6_1 and pages/6_2 prompt_stability_chart",
            color="#b45309",
        )
        attachments.append((p, "Figure A.11 - Prompt missing-tone benchmark", "Compares prompt templates by how often tone is missing, supporting RQ6 and Chapter 5 limitations."))

    ontology = load_csv(REV / "ontology_coverage.csv")
    if not ontology.empty and {"mapped_to_ontology", "records"}.issubset(ontology.columns):
        plot = ontology.copy()
        plot["records"] = pd.to_numeric(plot["records"], errors="coerce").fillna(0)
        mapped = plot.groupby(plot["mapped_to_ontology"].astype(str))["records"].sum().reset_index()
        mapped["label"] = mapped["mapped_to_ontology"].map({"True": "Mapped to ontology", "False": "Novel / unmapped"})
        p = draw_bar_chart(
            GRAPH_DIR / "docx_ontology_mapped_vs_unmapped.png",
            "Ontology Coverage Benchmark",
            zip(mapped["label"], mapped["records"]),
            subtitle="Source: pages/6_2 ontology_chart and Chapter 5 interpretation",
            color="#4f9d78",
        )
        attachments.append((p, "Figure A.12 - Ontology mapped vs novel aspects", "Shows whether ESG aspects are covered by existing ontology paths or represent Indonesian-specific vocabulary."))

    return attachments


def metric_summary() -> dict[str, str]:
    tone = load_csv(VIS / "tone_records_flat.csv")
    ocr = load_csv(REV / "ocr_processing_summary.csv")
    agreement = load_csv(REV / "climatebert_proxy_agreement_summary.csv")
    ontology = load_csv(REV / "ontology_coverage.csv")
    model = load_csv(REV / "model_stability_summary.csv")
    prompt = load_csv(REV / "prompt_stability_summary.csv")
    docs = max(tone["target_doc"].nunique() if not tone.empty and "target_doc" in tone.columns else 0, len(ocr))
    prompts = max(tone["prompt"].nunique() if not tone.empty and "prompt" in tone.columns else 0, len(prompt))
    pages = pd.to_numeric(ocr.get("pages", pd.Series(dtype=float)), errors="coerce").sum()
    mapped = int(ontology.get("mapped_to_ontology", pd.Series(dtype=bool)).astype(bool).sum()) if not ontology.empty and "mapped_to_ontology" in ontology.columns else 0
    kappa = pd.to_numeric(agreement.get("cohen_kappa", pd.Series(dtype=float)), errors="coerce")
    pct = pd.to_numeric(agreement.get("percent_agreement", pd.Series(dtype=float)), errors="coerce")
    return {
        "Tone records": f"{len(tone):,}",
        "Source documents": f"{docs:,}",
        "OCR documents": f"{len(ocr):,}",
        "OCR pages": f"{int(pages):,}",
        "Prompt templates": f"{prompts:,}",
        "Model rows": f"{len(model):,}",
        "ClimateBERT/proxy agreement": f"{pct.iloc[0]:.1%}" if not pct.empty and pd.notna(pct.iloc[0]) else "n/a",
        "Cohen kappa": f"{kappa.iloc[0]:.3f}" if not kappa.empty and pd.notna(kappa.iloc[0]) else "n/a",
        "Ontology mapped rows": f"{mapped}/{len(ontology):,}" if len(ontology) else "n/a",
    }


def add_caption(document: Document, caption: str, note: str = "") -> None:
    p = document.add_paragraph()
    run = p.add_run(caption)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(31, 41, 55)
    if note:
        p.add_run(f". {note}").font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)


def add_manual_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(23, 63, 66)
    run.font.size = Pt(18 if level == 1 else 14 if level == 2 else 12)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr[idx].text = text
        for paragraph in hdr[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = str(text)
    document.add_paragraph()


def add_graph(document: Document, path: Path, caption: str, note: str) -> None:
    if not path.exists():
        return
    add_caption(document, caption, note)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.4))
    paragraph.paragraph_format.space_after = Pt(12)


def update_document() -> Path:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)
    shutil.copy2(SOURCE_DOCX, OUTPUT_DOCX)
    document = Document(str(OUTPUT_DOCX))

    document.add_page_break()
    add_manual_heading(document, "Appendix A. Streamlit Graph Attachments and Chapter 4-6 Integration", level=1)
    document.add_paragraph(
        "This appendix updates the Chapter 4-6 structure and benchmark document with graph attachments "
        "and live evidence mappings from the Streamlit pages: pages/6_1_Chapter_4_Implementation_Results.py, "
        "pages/6_2_Chapter_5_Discussion.py, and pages/6_3_Chapter_6_Conclusion.py. The aim is to make the Word "
        "document usable alongside the Streamlit application: Chapter 4 carries implementation results, Chapter 5 "
        "interprets validation and limitations, and Chapter 6 turns those results into contributions and future work."
    )

    add_manual_heading(document, "A.1 Live Evidence Snapshot", level=2)
    metrics = metric_summary()
    add_table(document, ["Metric", "Current value", "Primary Streamlit source"], [[k, v, "6_1 / 6_2 / 6_3 shared data bundle"] for k, v in metrics.items()])

    add_manual_heading(document, "A.2 Streamlit Page to Thesis Chapter Mapping", level=2)
    add_table(
        document,
        ["Streamlit page", "Thesis role", "Graphs / evidence inserted"],
        [
            [
                "pages/6_1_Chapter_4_Implementation_Results.py",
                "Chapter 4 implementation and empirical results for RQ1-RQ6.",
                "PDF x prompt evidence, tone distribution, ESG by tone, ClimateBERT crosstab, artifact/model/prompt stability.",
            ],
            [
                "pages/6_2_Chapter_5_Discussion.py",
                "Chapter 5 interpretation, construct validity, limitations, and diagnostic discussion.",
                "Agreement metrics, ontology coverage, prompt/model sensitivity, failure-mode evidence.",
            ],
            [
                "pages/6_3_Chapter_6_Conclusion.py",
                "Chapter 6 contribution summary, research-question answers, and future work.",
                "Artifact inventory, workflow coverage, model/prompt benchmarks, contribution evidence.",
            ],
        ],
    )

    add_manual_heading(document, "A.3 Attached Graph Register", level=2)
    static_graphs = [
        (VIS / "tone_distribution.png", "Figure A.1 - Tone distribution", "Supports RQ2 and Chapter 4 by showing disclosure tone balance."),
        (VIS / "esg_by_tone.png", "Figure A.2 - ESG by tone", "Shows how ESG pillar claims vary by tone category."),
        (VIS / "aspect_by_tone_heatmap.png", "Figure A.3 - Aspect by tone heatmap", "Links aspect-level ABSA to tone patterns."),
        (VIS / "climatebert_label_by_tone.png", "Figure A.4 - Tone by ClimateBERT label", "Supports RQ3 construct comparison."),
        (VIS / "climatebert_remote_top_scores.png", "Figure A.5 - Top-scoring ClimateBERT records", "Provides examples for Chapter 5 interpretation."),
        (VIS / "streamlit_outputs" / "01_overview.png", "Figure A.6 - Streamlit overview", "Summarises the dashboard evidence layer."),
        (VIS / "streamlit_outputs" / "02_per_rq_evidence.png", "Figure A.7 - Per-RQ evidence", "Connects research questions to available outputs."),
        (VIS / "streamlit_outputs" / "04_benchmarks.png", "Figure A.8 - Benchmark plan", "Supports model, prompt, and validation benchmark framing."),
        (VIS / "streamlit_outputs" / "07_evidence_matrix.png", "Figure A.9 - Evidence matrix", "Connects thesis claims to result artifacts."),
    ]
    generated_graphs = create_generated_graphs()
    for path, caption, note in static_graphs + generated_graphs:
        add_graph(document, path, caption, note)

    add_manual_heading(document, "A.4 Chapter-Level Insert Notes", level=2)
    add_table(
        document,
        ["Chapter section", "Recommended insertion point", "Narrative update"],
        [
            [
                "4.2 RQ1-RQ2",
                "After the existing RQ1/RQ2 result paragraphs.",
                "Insert Figures A.1-A.3 to show tone, ESG pillar, and aspect evidence derived from pages/6_1.",
            ],
            [
                "4.3 RQ3-RQ4",
                "After the ClimateBERT and diagnostics paragraphs.",
                "Insert Figures A.4-A.5 and A.12 to connect ClimateBERT divergence and ontology coverage to Chapter 5.",
            ],
            [
                "4.4 RQ5-RQ6",
                "After the reproducibility and stability benchmark section.",
                "Insert Figures A.8, A.9, A.10, and A.11 to document benchmarking and artifact traceability.",
            ],
            [
                "5 Discussion",
                "At the beginning of the limitations and construct-validity discussion.",
                "Use the agreement, ontology, and prompt/model sensitivity graphs as evidence rather than treating limitations as prose-only claims.",
            ],
            [
                "6 Conclusion",
                "In contributions and future-work subsections.",
                "Use artifact inventory and benchmark graphics to define concrete future work: OCR quality metrics, more models, repeated runs, and ontology extension.",
            ],
        ],
    )

    add_manual_heading(document, "A.5 Benchmark Checklist Still Needed", level=2)
    add_table(
        document,
        ["Benchmark", "Why needed", "Target artifact"],
        [
            ["OCR quality", "Current page/document coverage exists, but CER/WER is not yet measured.", "ocr_quality_by_page.csv"],
            ["Human annotation agreement", "Single-annotator labels need inter-annotator reliability for stronger Chapter 5 validity.", "human_agreement_summary.csv"],
            ["Repeated LLM runs", "Model/prompt stability should include repeated runs and confidence intervals.", "model_prompt_repeated_run_ci.csv"],
            ["ClimateBERT baseline", "Tone-vs-ClimateBERT should be compared to majority and human-labelled baselines.", "climatebert_baseline_comparison.csv"],
            ["Ontology extension", "Unmapped Indonesian ESG aspects should be formalised as a vocabulary extension.", "indonesian_esg_ontology_extension.csv"],
        ],
    )

    document.save(str(OUTPUT_DOCX))
    return OUTPUT_DOCX


if __name__ == "__main__":
    out = update_document()
    print(out)
