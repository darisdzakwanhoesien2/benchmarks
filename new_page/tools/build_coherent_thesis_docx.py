from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "thesis_paper_esg_absa_combined_coherent.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, True)
        set_cell_shading(cell, "EAF2F8")
        if widths:
            cell.width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_callout(doc: Document, title: str, body: str, fill: str = "F4F9F4") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(37, 91, 65)
    p.add_run("\n" + body)
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Heading 1", 18, "1F4E5F"),
        ("Heading 2", 14, "245B41"),
        ("Heading 3", 11.5, "3E5F73"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(5)


def p(doc: Document, text: str, style: str | None = None, bold_lead: str | None = None) -> None:
    para = doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        r = para.add_run(bold_lead)
        r.bold = True
        para.add_run(text[len(bold_lead):])
    else:
        para.add_run(text)


def read_metrics() -> dict:
    metrics = {
        "dashboard_records": "332",
        "t2_rows": "2,074",
        "ocr_docs": "23",
        "artifacts": "40",
        "climatebert_agreement": "83.7%",
        "kappa": "0.645",
    }
    path = ROOT / "results/thesis_workflow_dashboard/dashboard_metrics.json"
    if path.exists():
        import json

        data = json.loads(path.read_text())
        metrics.update(
            {
                "dashboard_records": f"{int(data.get('tone_records', 332)):,}",
                "t2_rows": f"{int(data.get('t2_rows', 2074)):,}",
                "ocr_docs": f"{int(data.get('ocr_docs', 23)):,}",
                "artifacts": f"{int(data.get('artifacts', 40)):,}",
                "climatebert_agreement": f"{float(data.get('climatebert_percent_agreement', 0.8373493975903614))*100:.1f}%",
                "kappa": f"{float(data.get('climatebert_cohen_kappa', 0.6451446894422231)):.3f}",
            }
        )
    return metrics


def read_stability_rows() -> tuple[list[list[str]], list[list[str]]]:
    model_rows: list[list[str]] = []
    prompt_rows: list[list[str]] = []
    model_path = ROOT / "results/thesis_workflow_dashboard/model_stability_summary.csv"
    prompt_path = ROOT / "results/thesis_workflow_dashboard/prompt_stability_summary.csv"
    if model_path.exists():
        df = pd.read_csv(model_path)
        for _, row in df.head(6).iterrows():
            model_rows.append(
                [
                    str(row.get("model", "")),
                    str(int(row.get("runs", 0))),
                    f"{float(row.get('json_parse_success_rate', 0)):.3f}",
                    f"{float(row.get('avg_records', 0)):.2f}",
                    f"{float(row.get('missing_tone_rate', 0)):.3f}",
                ]
            )
    if prompt_path.exists():
        df = pd.read_csv(prompt_path)
        for _, row in df.head(8).iterrows():
            prompt_rows.append(
                [
                    str(row.get("prompt", "")),
                    str(int(row.get("runs", 0))),
                    f"{float(row.get('json_parse_success_rate', 0)):.3f}",
                    f"{float(row.get('missing_tone_rate', 0)):.3f}",
                    f"{float(row.get('field_completion_rate', 0)):.3f}",
                ]
            )
    return model_rows, prompt_rows


def read_analysis_rows() -> list[list[str]]:
    path = ROOT / "research_references/notes/Final Analysis - Table Description.csv"
    if not path.exists():
        return []
    df = pd.read_csv(path)
    wanted = ["A.1", "A.2", "A.3", "A.4", "A.10", "A.12", "A.30", "A.31", "A.33", "A.36"]
    rows = []
    for _, row in df[df["figure"].isin(wanted)].iterrows():
        explanation = str(row.get("Explanations", ""))
        assessment = str(row.get("Assessment", ""))
        rows.append(
            [
                str(row.get("figure", "")),
                str(row.get("title", "")),
                str(row.get("rq", "")),
                assessment.replace("Assessment:", "").strip()[:95],
                explanation[:180] + ("..." if len(explanation) > 180 else ""),
            ]
        )
    return rows


def build_doc() -> None:
    metrics = read_metrics()
    model_rows, prompt_rows = read_stability_rows()
    analysis_rows = read_analysis_rows()

    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Toward an Executable ESG ABSA Framework for Indonesian Sustainability Reports")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(31, 78, 95)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Coherent thesis paper edition with strengthened analysis and evidence hierarchy")
    r.italic = True
    r.font.size = Pt(11)

    p(doc, "Author: Daris Dzakwan Hoesien")
    p(doc, "Draft status: revised coherent edition")
    p(doc, "Date: 2026-05-22")

    add_callout(
        doc,
        "Editorial note",
        "This edition reorganizes the previous combined draft into a single coherent thesis argument. It separates the 332-record dashboard snapshot from the larger action-plan and silver-label evidence, removes duplicated section flow, and strengthens the interpretation of tone, ClimateBERT divergence, ontology depth, and reliability.",
        "EAF2F8",
    )

    doc.add_heading("Abstract", level=1)
    p(
        doc,
        "Sustainability reports contain dense qualitative and quantitative disclosures about environmental, social, and governance (ESG) activities, but their narrative structure makes them difficult to compare at the level of individual claims. This thesis proposes an executable ESG Aspect-Based Sentiment Analysis (ABSA) framework for Indonesian sustainability reports. The framework transforms PDF reports into structured ESG evidence records by combining OCR, page-aware batching, large language model extraction, tone-aware ABSA labeling, ClimateBERT comparison, ontology mapping, and Streamlit-based evidence dashboards.",
    )
    p(
        doc,
        f"The current dashboard evidence layer contains {metrics['dashboard_records']} structured tone records, {metrics['t2_rows']} T2 rows, {metrics['ocr_docs']} OCR documents, and {metrics['artifacts']} tracked artifacts. ClimateBERT proxy comparison reports {metrics['climatebert_agreement']} agreement and Cohen's kappa of {metrics['kappa']}. The larger action-plan and silver-label layer extends the validation scope beyond this dashboard sample, but it must be reported as a separate evidence tier. The main analytical contribution is therefore not a single classifier score; it is an auditable workflow for distinguishing ESG commitments, actions, outcomes, neutral disclosures, extraction failures, and ontology gaps.",
    )
    p(
        doc,
        "Keywords: ESG disclosure, aspect-based sentiment analysis, ClimateBERT, Indonesian sustainability reports, LLM extraction, ontology mapping, greenwashing diagnostics, reproducible dashboards.",
    )

    doc.add_heading("1. Introduction", level=1)
    doc.add_heading("1.1 Motivation", level=2)
    p(
        doc,
        "ESG disclosure analysis is increasingly important for investors, regulators, companies, and researchers. Yet sustainability reports are not simple sentiment documents. They combine regulatory disclosure, strategic narrative, performance claims, future commitments, operational descriptions, and sometimes promotional language. In the Indonesian setting, this complexity is intensified by bilingual reporting, OJK disclosure obligations, sector-specific terminology, and inconsistent PDF formatting.",
    )
    p(
        doc,
        "A document-level ESG score cannot explain whether a company is promising future action, reporting completed action, disclosing measurable outcome, or merely describing a governance process. This motivates a record-level ABSA framework where each disclosure unit is classified by aspect, ESG pillar, sentiment, tone, model provenance, prompt provenance, and source document.",
    )

    doc.add_heading("1.2 Research Problem", level=2)
    p(
        doc,
        "The research problem is how to transform Indonesian sustainability-report PDFs into structured, validated, and reproducible ESG evidence records. The pipeline must preserve provenance, support multi-label ESG categories, distinguish tone from sentiment, handle OCR and LLM failures, compare against ClimateBERT-style climate labels, and export the result into thesis-ready dashboards and graph formats.",
    )

    doc.add_heading("1.3 Research Questions", level=2)
    add_table(
        doc,
        ["RQ", "Question", "Primary evidence"],
        [
            ["RQ1", "How can PDF reports be transformed into structured ESG records while preserving provenance?", "OCR outputs, page audits, extracted ESG records"],
            ["RQ2", "How should ESG disclosures be categorized by aspect, ESG pillar, sentiment, and tone?", "Tone distribution, ESG-by-tone, aspect-by-tone heatmap"],
            ["RQ3", "How do ESG tone labels differ from ClimateBERT-style labels?", "Tone-by-ClimateBERT crosstab, agreement, kappa"],
            ["RQ4", "What failures and ontology gaps appear in the extraction pipeline?", "Failure modes, ontology coverage, aspect networks"],
            ["RQ5", "How can the workflow be made reproducible and auditable?", "Job folders, configs, artifact lineage, Streamlit pages"],
            ["RQ6", "How stable are outputs across prompts, models, and providers?", "Model stability and prompt stability summaries"],
        ],
        [0.6, 4.4, 2.2],
    )

    doc.add_heading("1.4 Contributions", level=2)
    add_numbered(
        doc,
        [
            "An end-to-end OCR-to-record ESG processing pipeline.",
            "A tone-aware ESG ABSA schema that separates sentiment from commitment/action/outcome tone.",
            "A ClimateBERT comparison framework that treats climate labels as construct comparison rather than final truth.",
            "A diagnostics layer for parse errors, schema drift, missing fields, OCR loss, and ontology gaps.",
            "An ontology and semantic export pathway for RDF, OWL, Neo4j, and future GraphRAG use.",
            "A Streamlit evidence layer that links research questions, backing tables, graphs, and thesis chapter claims.",
        ],
    )

    doc.add_heading("2. Related Work And Research Gap", level=1)
    p(
        doc,
        "The related work can be organized into four families: lexicon-based ESG scoring, classical machine learning and topic modeling, transformer/BERT-based domain models, and LLM-based extraction. Lexicon and classical methods provide transparency but often remain document-level or keyword-level. Transformer models such as ClimateBERT improve domain specificity but do not directly solve the broader ESG tone-taxonomy problem. LLM-based systems can produce richer structured outputs but require stronger reliability controls.",
    )
    p(
        doc,
        "The research gap is therefore integrative. Existing work rarely combines sentence-level ESG evidence extraction, bilingual Indonesian reporting context, tone taxonomy, ClimateBERT comparison, ontology mapping, prompt/model stability, and reproducible dashboard evidence in one pipeline.",
    )
    add_callout(
        doc,
        "Coherence improvement",
        "This paper positions ClimateBERT as a validator and comparison layer, not as a direct replacement for ESG ABSA. This prevents the analysis from treating disagreement as simple error when it may be evidence of construct divergence.",
        "FFF7E6",
    )

    doc.add_heading("3. Methodology", level=1)
    doc.add_heading("3.1 Evidence Tiers", level=2)
    p(
        doc,
        "The previous combined draft mixed several corpus sizes without enough explanation. This edition separates the evidence into tiers. The 332-record dashboard table is the compact analysis snapshot used for many visualizations. The 5,444-row action-plan and silver-label layer is the broader validation and annotation scaffold. The T2 and ground_truth.py outputs provide additional pipeline-specific evidence. These tiers should not be collapsed into one denominator unless the same unit of analysis is being measured.",
    )
    add_table(
        doc,
        ["Evidence tier", "Typical size", "Primary use", "Caution"],
        [
            ["Dashboard snapshot", metrics["dashboard_records"], "Chapter 4 figures and core ABSA visualizations", "Small enough for visual explanation; not always the full corpus"],
            ["T2 pipeline layer", metrics["t2_rows"], "Rule/hybrid tone and sentiment diagnostics", "Contains classifier-specific abstention behavior"],
            ["Action-plan / silver-label layer", "5,444 rows reported in notes", "Annotation completion, large-scale validation scaffold", "Must be separated from extracted-record counts"],
            ["Ground truth candidate layer", "large candidate pool", "Stress-testing local ClimateBERT and T1/T2 processing", "Failures may reflect domain/model mismatch and batching limits"],
        ],
        [1.7, 1.2, 2.6, 2.3],
    )

    doc.add_heading("3.2 Pipeline Design", level=2)
    add_numbered(
        doc,
        [
            "PDF ingestion and OCR convert sustainability reports into page-aware text while preserving document provenance.",
            "Page batching prepares text units suitable for LLM context windows and downstream record tracing.",
            "LLM extraction generates structured ESG records using prompt templates and model/provider metadata.",
            "ABSA normalization maps extracted text into aspect, ESG pillar, sentiment, and tone fields.",
            "ClimateBERT comparison provides external climate-oriented labels for construct comparison.",
            "Ground-truth workbench and silver-label tables support validation and review queues.",
            "Ontology mapping separates mapped aspects, shallow mappings, and novel Indonesian ESG vocabulary candidates.",
            "Streamlit dashboards and semantic exports translate the evidence layer into figures, tables, RDF, OWL, Neo4j, and chapter claims.",
        ],
    )

    doc.add_heading("3.3 Why Tone Is Separated From Sentiment", level=2)
    p(
        doc,
        "The central methodological choice is to distinguish sentiment from disclosure tone. Sentiment asks whether text is positive, neutral, or negative. Tone asks what kind of ESG claim the text is making: commitment, action, outcome, none, missing, or other. This distinction is necessary because ESG reports frequently use positive or neutral language for claims that differ substantially in accountability. A future commitment and a measured outcome should not be treated as the same evidence.",
    )

    doc.add_heading("4. Results And Analysis", level=1)
    doc.add_heading("4.1 Pipeline Coverage", level=2)
    p(
        doc,
        f"The current dashboard layer contains {metrics['dashboard_records']} structured tone records from {metrics['ocr_docs']} OCR documents, with {metrics['artifacts']} tracked artifacts. This establishes that the PDF-to-record pipeline is operational. However, coverage should be interpreted at the correct evidence tier: the dashboard snapshot supports interpretive visualization, while the larger action-plan and silver-label datasets support annotation completeness and broader validation claims.",
    )

    doc.add_heading("4.2 Tone Distribution And ESG Pillar Analysis", level=2)
    p(
        doc,
        "The notes report that the full tone distribution is dominated by the none category. This is analytically important rather than merely a pipeline weakness. It suggests that a substantial portion of sustainability-report text is descriptive or procedural rather than performative. A coherent interpretation must therefore avoid equating every ESG mention with a commitment, action, or outcome.",
    )
    p(
        doc,
        "The ESG-by-tone analysis also shows that tone is not evenly distributed across pillars. Governance can dominate in Indonesian reports because OJK-style reporting requires structured narrative about governance mechanisms, board oversight, compliance, and anti-corruption policies. This may invert expectations from English-language ESG corpora where environmental action/commitment language is often more dominant.",
    )

    doc.add_heading("4.3 Aspect-Tone Signatures", level=2)
    p(
        doc,
        "The aspect-by-tone heatmap is the strongest validation of the tone taxonomy. Aspects such as sustainability commitment naturally concentrate in commitment tone, while ESG performance concentrates in outcome tone. This pattern indicates that the model is not assigning labels randomly; it is capturing semantically coherent aspect-tone relationships.",
    )
    p(
        doc,
        "The same analysis supports greenwashing diagnostics. Aspects with high commitment counts and low outcome counts are potential candidates for commitment-outcome asymmetry. Governance and anti-corruption topics are especially important because firms may pledge integrity or compliance without reporting operationalized outcomes.",
    )

    doc.add_heading("4.4 ClimateBERT Comparison", level=2)
    p(
        doc,
        f"The ClimateBERT comparison reports {metrics['climatebert_agreement']} raw agreement and kappa={metrics['kappa']}. This should be interpreted as moderate construct alignment, not as final human validation. ClimateBERT is trained for climate-related language, whereas the ESG ABSA pipeline separates tone categories across environmental, social, and governance disclosures. Disagreement is therefore not automatically a failure; it can reveal where climate relevance and disclosure accountability diverge.",
    )
    add_callout(
        doc,
        "Key RQ3 interpretation",
        "The novel finding is not that the pipeline perfectly matches ClimateBERT. The more interesting finding is that the systems diverge in interpretable ways because they operationalize different constructs.",
        "F4F9F4",
    )

    doc.add_heading("4.5 Ontology Coverage", level=2)
    p(
        doc,
        "The ontology results need a depth-versus-breadth interpretation. A record may be technically mapped while still landing on a shallow category such as Governance -> General or a ClimateBERT placeholder. This means mapped/unmapped counts alone are insufficient. Regulatory traceability requires deeper links into GRI, SASB, TCFD, OJK, or thesis-specific Indonesian ESG ontology nodes.",
    )
    p(
        doc,
        "Unmapped or shallowly mapped aspects are not only limitations. They are also evidence for the thesis contribution: Indonesian ESG reporting contains vocabulary and disclosure patterns that require ontology extension.",
    )

    doc.add_heading("4.6 Model And Prompt Stability", level=2)
    if model_rows:
        add_table(
            doc,
            ["Model", "Runs", "Parse success", "Avg. records", "Missing-tone"],
            model_rows,
            [2.8, 0.6, 0.9, 0.9, 0.9],
        )
    if prompt_rows:
        add_table(
            doc,
            ["Prompt", "Runs", "Parse success", "Missing-tone", "Field completion"],
            prompt_rows,
            [2.7, 0.5, 0.9, 0.9, 1.0],
        )
    p(
        doc,
        "The stability evidence shows that successful JSON parsing is not enough. A run can parse correctly while still producing missing tone fields, shallow records, or incomplete ABSA dimensions. Therefore, model evaluation should report parse success, missing-tone rate, schema drift, and field completion together.",
    )

    doc.add_heading("4.7 Figure-Level Interpretation", level=2)
    if analysis_rows:
        add_table(
            doc,
            ["Figure", "Title", "RQ", "Assessment", "Interpretive use"],
            analysis_rows,
            [0.55, 1.65, 0.75, 1.25, 3.0],
        )

    doc.add_heading("5. Discussion", level=1)
    doc.add_heading("5.1 Coherent Interpretation Of The Main Findings", level=2)
    add_bullets(
        doc,
        [
            "The dominance of none or descriptive tone should be read as a corpus property and calibration question, not only as model error.",
            "Governance-heavy tone patterns are plausible in Indonesia because regulatory reporting obligations emphasize governance disclosure.",
            "Aspect-tone concentration validates the taxonomy when dominant tones align with semantic expectations.",
            "ClimateBERT disagreement supports construct validity analysis because climate relevance and ESG disclosure tone are different tasks.",
            "Ontology mapping should be judged by regulatory specificity, not only by whether an aspect receives any path.",
            "Prompt/model stability must include missing-field and schema-drift measures, not only parse success.",
        ],
    )

    doc.add_heading("5.2 Limitations", level=2)
    add_bullets(
        doc,
        [
            "OCR quality has not yet been fully measured with CER/WER across the required sampled pages.",
            "The 332-record dashboard snapshot and 5,444-row silver/action-plan scaffold represent different units of analysis.",
            "The kappa value is a proxy agreement between automated systems, not completed human inter-annotator agreement.",
            "ClimateBERT is climate-oriented and English-heavy, so it is not a full ESG benchmark for Indonesian bilingual reports.",
            "Ontology coverage currently includes shallow paths; deeper regulatory mapping remains a development priority.",
            "Cross-entity and temporal claims require stronger normalization and year metadata extraction.",
        ],
    )

    doc.add_heading("5.3 Stronger Thesis Claim", level=2)
    p(
        doc,
        "The strongest defensible claim is that the thesis demonstrates an executable, auditable ESG ABSA research pipeline, not that it has already solved fully supervised ESG classification. Its empirical contribution lies in showing how record-level tone, ClimateBERT divergence, ontology depth, prompt stability, and graph-ready artifacts can be integrated into one reproducible evidence system.",
    )

    doc.add_heading("6. Conclusion", level=1)
    p(
        doc,
        "This thesis proposes a reproducible ESG ABSA framework for Indonesian sustainability reports. The framework addresses a real methodological gap: ESG disclosure cannot be reliably interpreted through document-level sentiment alone. The record-level pipeline makes it possible to distinguish commitments, actions, outcomes, neutral descriptions, missing labels, and ontology gaps.",
    )
    p(
        doc,
        "The revised analysis clarifies the relationship between evidence tiers, strengthens the interpretation of ClimateBERT divergence, and reframes ontology gaps as both a limitation and a contribution. Future work should complete OCR quality measurement, expand human annotation, improve regulatory ontology depth, normalize cross-company comparisons, and connect the semantic exports to Neo4j and GraphRAG workflows.",
    )

    doc.add_heading("7. References", level=1)
    refs = [
        "McHugh, M. L. (2012). Interrater reliability: The kappa statistic. Biochemia Medica, 22(3), 276-282. https://pmc.ncbi.nlm.nih.gov/articles/PMC3900052/",
        "Webersinke, N., Kraus, M., Bingler, J. A., & Leippold, M. (2021). ClimateBERT: A pretrained language model for climate-related text. arXiv. https://arxiv.org/abs/2110.12010",
        "Zhang, W., Li, X., Deng, Y., Bing, L., & Lam, W. (2022). A survey on aspect-based sentiment analysis: Tasks, methods, and challenges. arXiv. https://arxiv.org/abs/2203.01054",
        "Islam, M. M. (2025). Towards Behaviour-Aware Multimodal Video Summarization: Integrating Visual, Audio, and Textual Cues for Human-Centric Content Analysis [Master's thesis, University of Oulu]. Local reference file: research_references/nbnfioulu-202506124422.pdf",
    ]
    for ref in refs:
        p(doc, ref)

    doc.add_heading("Appendix A. Artifact And Page Map", level=1)
    add_table(
        doc,
        ["Artifact / page", "Use in thesis"],
        [
            ["pages/3_0_Thesis_Action_Plan.py", "Live operational status, annotation progress, PDF-by-prompt matrix, migration controls"],
            ["pages/6_4_ch4-6.py", "Graph attachments, backing tables, chapter mapping, benchmark checklist"],
            ["pages/6_0_Thesis_Draft_Chapter_Integration_Mermaid.py", "Thesis spine, RQ evidence, validation loop, artifact lineage"],
            ["results/thesis_workflow_dashboard/tone_records_flat.csv", "Core record-level ESG ABSA table"],
            ["results/thesis_workflow_dashboard/model_stability_summary.csv", "Model stability and parse-success evidence"],
            ["results/thesis_workflow_dashboard/prompt_stability_summary.csv", "Prompt stability and missing-tone evidence"],
            ["results/thesis_workflow_dashboard/ontology_coverage.csv", "Mapped, unmapped, and shallow ontology paths"],
            ["results/semantic_exports/esg_thesis_graph.ttl", "RDF export for semantic ESG evidence"],
        ],
        [3.1, 4.0],
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
