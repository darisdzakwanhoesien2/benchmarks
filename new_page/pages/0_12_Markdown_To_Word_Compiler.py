from __future__ import annotations

from io import BytesIO
from pathlib import Path
import re

import pandas as pd
import streamlit as st
from _page_runtime_controls import apply_page_runtime_controls

try:
    from docx import Document
except ImportError:
    Document = None


st.set_page_config(page_title="Markdown to Word Compiler", layout="wide")
apply_page_runtime_controls(__file__)
st.title("Markdown to Word Compiler")
st.caption("Find Markdown files, combine selected files, preview the merged text, and download a Word document.")

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SEARCH_DIRS = ["documentation", "pages", "results", "."]


@st.cache_data(show_spinner=False)
def discover_markdown_files(base_dir: str) -> pd.DataFrame:
    base = Path(base_dir)
    files = sorted(base.rglob("*.md"))
    rows: list[dict[str, object]] = []
    for path in files:
        try:
            stat = path.stat()
        except OSError:
            continue
        rows.append(
            {
                "path": str(path.relative_to(base)),
                "name": path.name,
                "parent": str(path.parent.relative_to(base)) or ".",
                "size_kb": round(stat.st_size / 1024, 2),
                "modified": pd.to_datetime(stat.st_mtime, unit="s"),
                "abs_path": str(path),
            }
        )
    return pd.DataFrame(rows)


def normalize_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", value.strip())
    cleaned = re.sub(r"_+", "_", cleaned).strip("._")
    return cleaned or "markdown_bundle"


def read_markdown_text(path: str) -> str:
    return Path(path).read_text(encoding="utf-8", errors="ignore")


def path_matches_parent_filters(path_value: str, selected_parents: list[str]) -> bool:
    return any(
        path_value == parent or path_value.startswith(f"{parent}/")
        for parent in selected_parents
    )


def build_combined_markdown(selected_paths: list[str], add_file_headers: bool) -> str:
    parts: list[str] = []
    for selected_path in selected_paths:
        path = ROOT / selected_path
        content = read_markdown_text(str(path)).strip()
        if add_file_headers:
            parts.append(f"# {selected_path}\n")
        parts.append(content)
    return "\n\n---\n\n".join(part for part in parts if part)


def add_markdown_line(doc: Document, line: str, in_code_block: bool) -> None:
    stripped = line.rstrip()
    if not stripped:
        doc.add_paragraph("")
        return

    if in_code_block:
        doc.add_paragraph(stripped, style="No Spacing")
        return

    heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
    if heading_match:
        level = min(len(heading_match.group(1)), 4)
        doc.add_heading(heading_match.group(2).strip(), level=level)
        return

    bullet_match = re.match(r"^\s*[-*+]\s+(.*)$", stripped)
    if bullet_match:
        doc.add_paragraph(bullet_match.group(1).strip(), style="List Bullet")
        return

    numbered_match = re.match(r"^\s*\d+\.\s+(.*)$", stripped)
    if numbered_match:
        doc.add_paragraph(numbered_match.group(1).strip(), style="List Number")
        return

    doc.add_paragraph(stripped)


def build_docx_bytes(selected_paths: list[str], doc_title: str, add_file_headers: bool) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx is not installed.")

    doc = Document()
    doc.add_heading(doc_title, level=0)
    doc.add_paragraph(f"Compiled from {len(selected_paths)} Markdown files.")

    for index, selected_path in enumerate(selected_paths, start=1):
        if index > 1:
            doc.add_page_break()

        doc.add_heading(f"{index}. {selected_path}", level=1)
        content = read_markdown_text(str(ROOT / selected_path))

        in_code_block = False
        for raw_line in content.splitlines():
            if raw_line.strip().startswith("```"):
                in_code_block = not in_code_block
                continue
            add_markdown_line(doc, raw_line, in_code_block)

        if add_file_headers and not content.strip():
            doc.add_paragraph("(Empty file)")

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


catalog = discover_markdown_files(str(ROOT))
if catalog.empty:
    st.warning("No Markdown files were found in this repository.")
    st.stop()

with st.sidebar:
    st.header("Selection")
    parent_options = sorted(catalog["parent"].dropna().unique().tolist())
    selected_parents = st.multiselect(
        "Folders",
        parent_options,
        default=[parent for parent in parent_options if any(parent == d or parent.startswith(f"{d}/") for d in DEFAULT_SEARCH_DIRS)],
    )
    path_query = st.text_input("Path contains", value="")
    max_size_kb = st.number_input("Max size (KB)", min_value=0.0, value=float(catalog["size_kb"].max()), step=50.0)
    add_file_headers = st.checkbox("Insert file path headers into combined markdown preview", value=True)

filtered = catalog.copy()
if selected_parents:
    filtered = filtered[
        filtered["path"].apply(lambda value: path_matches_parent_filters(str(value), selected_parents))
    ]
filtered = filtered[filtered["size_kb"] <= max_size_kb]
if path_query.strip():
    filtered = filtered[filtered["path"].str.contains(path_query.strip(), case=False, na=False)]
filtered = filtered.sort_values(["parent", "path"]).reset_index(drop=True)

st.subheader("Markdown Inventory")
c1, c2, c3 = st.columns(3)
c1.metric("Markdown files", f"{len(catalog):,}")
c2.metric("Files after filters", f"{len(filtered):,}")
c3.metric("Filtered size", f"{filtered['size_kb'].sum():,.1f} KB")

st.dataframe(
    filtered[["path", "parent", "size_kb", "modified"]],
    use_container_width=True,
    height=320,
)

if filtered.empty:
    st.info("No Markdown files match the current filters.")
    st.stop()

default_selection = filtered["path"].head(10).tolist()
selected_paths = st.multiselect(
    "Select Markdown files to concatenate",
    filtered["path"].tolist(),
    default=default_selection,
)

if not selected_paths:
    st.info("Select at least one Markdown file to preview and export.")
    st.stop()

selected_df = filtered[filtered["path"].isin(selected_paths)].copy()
selected_df["selection_order"] = selected_df["path"].apply(lambda value: selected_paths.index(value))
selected_df = selected_df.sort_values("selection_order")
ordered_paths = selected_df["path"].tolist()

st.subheader("Selected Files")
st.dataframe(
    selected_df[["path", "size_kb", "modified"]],
    use_container_width=True,
    height=min(80 + len(selected_df) * 35, 360),
)

doc_title = st.text_input("Word document title", value="Markdown Compilation")
download_name = normalize_filename(st.text_input("Download filename", value="markdown_compilation")) + ".docx"

combined_markdown = build_combined_markdown(ordered_paths, add_file_headers=add_file_headers)

st.subheader("Combined Preview")
preview_chars = st.slider("Preview characters", min_value=1000, max_value=50000, value=12000, step=1000)
st.code(combined_markdown[:preview_chars], language="markdown")
if len(combined_markdown) > preview_chars:
    st.caption(f"Preview truncated to the first {preview_chars:,} characters.")

st.subheader("Download")
if Document is None:
    st.error("`python-docx` is not installed in the current environment. Install dependencies from `requirements.txt` to enable Word export.")
else:
    docx_bytes = build_docx_bytes(ordered_paths, doc_title=doc_title, add_file_headers=add_file_headers)
    st.download_button(
        "Download combined Word document",
        data=docx_bytes,
        file_name=download_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )
